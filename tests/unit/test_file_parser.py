"""Unit tests for modules.file_parser.FileParser."""

from __future__ import annotations

from pathlib import Path

import pytest

from core.service_container import ServiceContainer
from modules.file_parser import FileParser


@pytest.fixture
def parser(project_root: Path, tmp_path: Path) -> FileParser:
    """FileParser with production-like container."""
    container = ServiceContainer.create_production_container(
        app_config={
            "database_path": str(tmp_path / "autopilot.db"),
            "schema_path": str(project_root / "database" / "schema.sql"),
            "config_folder": str(project_root / "config"),
            "cache_folder": str(tmp_path / "cache"),
            "log_folder": str(tmp_path / "logs"),
            "ffmpeg_path": "ffmpeg",
        },
        project_root=project_root,
    )
    return FileParser(container)


@pytest.fixture
def sample_project(project_root: Path) -> Path:
    """Path to sample project fixtures."""
    return project_root / "tests" / "fixtures" / "sample_project"


class TestParseTxt:
    """TXT format parsing tests."""

    def test_parse_sample_txt(self, parser: FileParser, sample_project: Path) -> None:
        path = sample_project / "script" / "sample_script.txt"
        result = parser.parse_script(path)
        assert result["success"] is True
        data = result["data"]
        assert (
            data["project_settings"]["title"] == "The Dark History of the Black Death"
        )
        assert data["project_settings"]["genre"] == "dark_history"
        assert len(data["scenes"]) == 10
        assert data["scenes"][0]["image"] == "dark_castle_night.jpg"
        assert data["scenes"][0]["is_chapter_start"] is True
        assert data["scenes"][0]["chapter_title"] == "Introduction"
        assert data["scenes"][0]["dialogue"][0]["character"] == "NARRATOR"
        assert "1347" in data["scenes"][0]["dialogue"][0]["text"]
        voices = {v["character"]: v for v in data["voice_instructions"]}
        assert "NARRATOR" in voices
        assert voices["NARRATOR"]["engine"] == "piper"
        assert voices["NARRATOR"]["breathing"] is True

    def test_inline_pause_removed_from_text(
        self, parser: FileParser, sample_project: Path
    ) -> None:
        result = parser.parse_txt(sample_project / "script" / "sample_script.txt")
        first = result["data"]["scenes"][0]["dialogue"][0]["text"]
        assert "[PAUSE" not in first


class TestParseJsonCsv:
    """JSON and CSV parsing tests."""

    def test_parse_sample_json(self, parser: FileParser, sample_project: Path) -> None:
        result = parser.parse_json(sample_project / "script" / "sample_script.json")
        assert result["success"] is True
        data = result["data"]
        assert len(data["scenes"]) >= 2
        assert data["scenes"][0]["image"] == "dark_castle_night.jpg"
        assert data["voice_instructions"]

    def test_parse_sample_csv(self, parser: FileParser, sample_project: Path) -> None:
        result = parser.parse_csv(sample_project / "script" / "sample_script.csv")
        assert result["success"] is True
        data = result["data"]
        assert len(data["scenes"]) >= 2
        assert any(s["dialogue"] for s in data["scenes"])


class TestHelpers:
    """Helper method tests."""

    def test_parse_time_value(self, parser: FileParser) -> None:
        assert parser.parse_time_value("auto") is None
        assert parser.parse_time_value("90") == 90.0
        assert parser.parse_time_value("90s") == 90.0
        assert parser.parse_time_value("1:30") == 90.0
        assert parser.parse_time_value("00:01:30") == 90.0
        assert parser.parse_time_value("1m30s") == 90.0

    def test_normalize_names(self, parser: FileParser) -> None:
        assert parser.normalize_transition_name("cross_fade") == "crossfade"
        assert parser.normalize_transition_name("dip to black") == "dip_to_black"
        assert parser.normalize_animation_name("Ken Burns") == "ken_burns"
        assert parser.normalize_animation_name("zoom_in") == "slow_zoom_in"

    def test_detect_format(self, parser: FileParser, sample_project: Path) -> None:
        result = parser.detect_format(sample_project / "script" / "sample_script.txt")
        assert result["success"] is True
        assert result["data"]["format"] == "txt"


class TestImageMatching:
    """Filename exact + fuzzy matching."""

    def test_exact_match_all_sample_images(
        self, parser: FileParser, sample_project: Path
    ) -> None:
        txt = parser.parse_txt(sample_project / "script" / "sample_script.txt")
        result = parser.match_images(txt["data"]["scenes"], sample_project / "images")
        assert result["success"] is True
        assert result["data"]["total"] == 10
        assert result["data"]["exact"] == 10
        assert result["data"]["unmatched"] == 0

    def test_fuzzy_match(self, parser: FileParser, sample_project: Path) -> None:
        scenes = [{"id": "s1", "image": "dark_castle_nite.jpg"}]
        result = parser.match_images(scenes, sample_project / "images")
        assert result["success"] is True
        match = result["data"]["matches"][0]
        assert match["status"] in ("fuzzy", "exact")
        assert match["path"] is not None


class TestDocxPdf:
    """DOCX/PDF parsers when libraries available."""

    def test_parse_docx_paragraphs(self, parser: FileParser, tmp_path: Path) -> None:
        pytest.importorskip("docx")
        from docx import Document

        path = tmp_path / "script.docx"
        doc = Document()
        doc.add_paragraph("NARRATOR (dramatic): Ships arrived in Sicily.")
        doc.add_paragraph("HISTORIAN: Records confirm the spread.")
        doc.save(path)
        result = parser.parse_docx(path)
        assert result["success"] is True
        assert len(result["data"]["scenes"][0]["dialogue"]) >= 2

    def test_parse_pdf_text(self, parser: FileParser, tmp_path: Path) -> None:
        # PDF generation without extra deps is awkward; if reportlab missing, skip
        pytest.importorskip("pdfplumber")
        try:
            from reportlab.pdfgen import canvas  # type: ignore
        except ImportError:
            pytest.skip("reportlab not installed for PDF generation")
        path = tmp_path / "script.pdf"
        c = canvas.Canvas(str(path))
        c.drawString(100, 750, "//SCENE_START: scene_01")
        c.drawString(100, 730, "//IMAGE: dark_castle_night.jpg")
        c.drawString(100, 710, "[NARRATOR|dramatic]")
        c.drawString(100, 690, "The year was 1347.")
        c.drawString(100, 670, "//SCENE_END")
        c.save()
        result = parser.parse_pdf(path)
        assert result["success"] is True
        assert result["warnings"]


class TestClassifyAndZip:
    """Folder classification and ZIP extract."""

    def test_classify_sample_folder(
        self, parser: FileParser, sample_project: Path
    ) -> None:
        result = parser.classify_project_folder(sample_project)
        assert result["success"] is True
        assert result["data"]["scripts"]
        assert result["data"]["images"]
        assert result["data"]["audio"]

    def test_extract_zip(
        self, parser: FileParser, sample_project: Path, tmp_path: Path
    ) -> None:
        import zipfile

        zip_path = tmp_path / "project.zip"
        with zipfile.ZipFile(zip_path, "w") as archive:
            for file_path in sample_project.rglob("*"):
                if file_path.is_file():
                    archive.write(file_path, file_path.relative_to(sample_project))
        dest = tmp_path / "extracted"
        result = parser.extract_zip(zip_path, dest)
        assert result["success"] is True
        assert result["data"]["images"]
