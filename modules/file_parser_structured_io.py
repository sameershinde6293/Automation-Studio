"""Structured script file IO wrappers (JSON/CSV/DOCX/PDF)."""

from __future__ import annotations

import json
import re
import time
from pathlib import Path
from typing import List

from core.errors import ScriptParseError
from modules.file_parser_helpers import _elapsed_ms, _read_text


class StructuredScriptIOMixin:
    """File-format entry points and validation helpers."""

    def parse_json_file(self, file_path, make_response, log):
        """Parse JSON script file."""
        import time

        started = time.perf_counter()
        try:
            raw = json.loads(_read_text(Path(file_path)))
            data = self._map_json_to_parsed(raw)
            return make_response(True, data=data, duration_ms=_elapsed_ms(started))
        except json.JSONDecodeError as exc:
            return make_response(
                False,
                error=f"Invalid JSON script: {exc}",
                duration_ms=_elapsed_ms(started),
            )
        except Exception as exc:  # noqa: BLE001
            log.error("parse_json failed: %s", exc, exc_info=True)
            return make_response(
                False, error=str(exc), duration_ms=_elapsed_ms(started)
            )

    def parse_csv_file(self, file_path, make_response, log):
        """Parse CSV script file."""
        started = time.perf_counter()
        try:
            data = self._parse_csv_text(_read_text(Path(file_path)))
            return make_response(True, data=data, duration_ms=_elapsed_ms(started))
        except Exception as exc:  # noqa: BLE001
            log.error("parse_csv failed: %s", exc, exc_info=True)
            return make_response(
                False, error=str(exc), duration_ms=_elapsed_ms(started)
            )

    def parse_docx_file(self, file_path, make_response, log):
        """Parse DOCX script file."""
        started = time.perf_counter()
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return make_response(
                False,
                error="python-docx is not installed",
                duration_ms=_elapsed_ms(started),
            )
        try:
            doc = Document(str(file_path))
            if doc.tables:
                rows = [
                    [(cell.text or "").strip() for cell in row.cells]
                    for row in doc.tables[0].rows
                ]
                data = self._parse_table_rows(rows)
            else:
                paragraphs = [
                    p.text for p in doc.paragraphs if p.text and p.text.strip()
                ]
                data = self._parse_paragraph_script(paragraphs)
            return make_response(True, data=data, duration_ms=_elapsed_ms(started))
        except Exception as exc:  # noqa: BLE001
            log.error("parse_docx failed: %s", exc, exc_info=True)
            return make_response(
                False, error=str(exc), duration_ms=_elapsed_ms(started)
            )

    def parse_pdf_file(self, file_path, make_response, log, txt_engine):
        """Parse PDF script file."""
        started = time.perf_counter()
        try:
            import pdfplumber  # type: ignore
        except ImportError:
            return make_response(
                False,
                error="pdfplumber is not installed",
                duration_ms=_elapsed_ms(started),
            )
        try:
            tables: List = []
            texts: List[str] = []
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    tables.extend(page.extract_tables() or [])
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        texts.append(page_text)
            warnings = [
                "PDF parsing may be imperfect. If results are wrong use TXT or JSON format."
            ]
            if tables:
                cleaned = [[(c or "").strip() for c in row] for row in tables[0]]
                data = self._parse_table_rows(cleaned)
            else:
                data = txt_engine._parse_txt_content("\n".join(texts))
            return make_response(
                True, data=data, warnings=warnings, duration_ms=_elapsed_ms(started)
            )
        except Exception as exc:  # noqa: BLE001
            log.error("parse_pdf failed: %s", exc, exc_info=True)
            return make_response(
                False, error=str(exc), duration_ms=_elapsed_ms(started)
            )

    def validate_parsed_data(self, parsed_data, make_response):
        """Validate required fields in parsed script data."""
        errors = []
        scenes = parsed_data.get("scenes") or []
        if not scenes:
            errors.append("No scenes found")
        characters_used = set()
        for index, scene in enumerate(scenes, start=1):
            if not scene.get("image"):
                errors.append(f"Scene {index} missing image")
            dialogue = scene.get("dialogue") or []
            if not dialogue:
                errors.append(f"Scene {index} has no dialogue")
            for line in dialogue:
                if not line.get("character"):
                    errors.append(f"Scene {index} dialogue missing character")
                if not str(line.get("text", "")).strip():
                    errors.append(f"Scene {index} dialogue missing text")
                characters_used.add(str(line.get("character", "")).upper())
        voice_chars = {
            str(v.get("character", "")).upper()
            for v in parsed_data.get("voice_instructions") or []
        }
        for character in sorted(characters_used):
            if character and character not in voice_chars:
                errors.append(f"No voice instruction for character {character}")
        return make_response(True, {"valid": len(errors) == 0, "errors": errors})

    def parse_time_value(self, time_string: str):
        """Parse flexible time strings to seconds."""
        if time_string is None:
            return None
        value = str(time_string).strip().lower()
        if not value or value == "auto":
            return None
        if re.fullmatch(r"\d+(\.\d+)?", value):
            return float(value)
        if value.endswith("s") and re.fullmatch(r"\d+(\.\d+)?s", value):
            return float(value[:-1])
        if value.endswith("m") and re.fullmatch(r"\d+(\.\d+)?m", value):
            return float(value[:-1]) * 60.0
        match = re.fullmatch(r"(?:(\d+)m)?(?:(\d+(?:\.\d+)?)s)?", value)
        if match and (match.group(1) or match.group(2)):
            minutes = float(match.group(1) or 0)
            seconds = float(match.group(2) or 0)
            return minutes * 60.0 + seconds
        if ":" in value:
            parts = [float(p) for p in value.split(":")]
            if len(parts) == 3:
                return parts[0] * 3600 + parts[1] * 60 + parts[2]
            if len(parts) == 2:
                return parts[0] * 60 + parts[1]
        raise ScriptParseError(f"Unrecognized time value: {time_string}")
